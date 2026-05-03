import os
import json
import threading

import customtkinter as ctk
from tkinter import filedialog, messagebox
import cv2
from scenedetect import open_video, SceneManager
from scenedetect.detectors import ContentDetector
import whisper
from docx import Document
from docx.shared import Inches, Pt


ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

WHISPER_MODELS = {
    "tiny   — быстро, менее точно": "tiny",
    "base   — баланс скорости и качества": "base",
    "small  — хорошее качество": "small",
    "medium — высокое качество": "medium",
    "large  — лучшее качество, медленно": "large",
}

FALLBACK_INTERVAL_SEC = 30


def _sensitivity_to_threshold(s: float) -> float:
    return round(32.0 / s, 1)


def _fmt_time(sec: float) -> str:
    h = int(sec // 3600)
    m = int((sec % 3600) // 60)
    s = int(sec % 60)
    if h > 0:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def _text_for_range(segments: list, t_from: float, t_to: float) -> str:
    parts = [s["text"].strip() for s in segments if t_from <= s["start"] < t_to]
    return " ".join(parts)


def _generate_word(out_dir: str, video_name: str, segments: list, screenshots: list) -> str:
    doc = Document()
    doc.add_heading(video_name, level=1)

    shots = sorted(screenshots, key=lambda x: x["time_sec"])
    end_time = segments[-1]["end"] if segments else 0.0

    # Текст до первого скриншота
    first_t = shots[0]["time_sec"] if shots else end_time + 1
    pre_text = _text_for_range(segments, 0.0, first_t)
    if pre_text:
        p = doc.add_paragraph()
        p.add_run("[00:00] ").bold = True
        p.add_run(pre_text)

    # Для каждого скриншота: картинка → метка → текст до следующего скриншота
    for i, shot in enumerate(shots):
        t_start = shot["time_sec"]
        t_end = shots[i + 1]["time_sec"] if i + 1 < len(shots) else end_time + 1

        doc.add_paragraph()  # отступ

        # Скриншот
        img_path = os.path.join(out_dir, shot["filename"])
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = 1  # CENTER
            p.add_run().add_picture(img_path, width=Inches(5.5))

        # Временна́я метка под картинкой
        cap = doc.add_paragraph(f"[{_fmt_time(t_start)}]  •  Скриншот {shot['index']}")
        cap.alignment = 1
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].bold = True

        # Текст этого раздела
        text = _text_for_range(segments, t_start, t_end)
        if text:
            p = doc.add_paragraph()
            p.add_run(f"[{_fmt_time(t_start)}] ").bold = True
            p.add_run(text)

    docx_path = os.path.join(out_dir, f"{video_name}.docx")
    doc.save(docx_path)
    return docx_path


class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Транскрибатор")
        self.geometry("700x620")
        self.resizable(False, False)
        self.video_path = None
        self._build_ui()

    # ── UI ──────────────────────────────────────────────────────────────────

    def _build_ui(self):
        ctk.CTkLabel(self, text="Транскрибатор видео", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(20, 5))

        file_frame = ctk.CTkFrame(self)
        file_frame.pack(fill="x", padx=20, pady=8)
        ctk.CTkButton(file_frame, text="Выбрать видео", width=140, command=self.select_video).pack(side="left", padx=10, pady=10)
        self.path_label = ctk.CTkLabel(file_frame, text="Файл не выбран", anchor="w", width=480)
        self.path_label.pack(side="left", padx=5)

        model_frame = ctk.CTkFrame(self)
        model_frame.pack(fill="x", padx=20, pady=4)
        ctk.CTkLabel(model_frame, text="Модель Whisper:", width=140).pack(side="left", padx=10, pady=10)
        self.model_var = ctk.StringVar(value=list(WHISPER_MODELS.keys())[1])
        ctk.CTkOptionMenu(model_frame, variable=self.model_var, values=list(WHISPER_MODELS.keys()), width=360).pack(side="left", padx=5)

        sens_frame = ctk.CTkFrame(self)
        sens_frame.pack(fill="x", padx=20, pady=4)
        ctk.CTkLabel(sens_frame, text="Скриншотов:", width=140).pack(side="left", padx=10, pady=10)
        ctk.CTkLabel(sens_frame, text="мало").pack(side="left")
        self.sens_slider = ctk.CTkSlider(sens_frame, from_=1, to=10, number_of_steps=9, width=260,
                                         command=self._on_slider)
        self.sens_slider.set(5)
        self.sens_slider.pack(side="left", padx=8)
        ctk.CTkLabel(sens_frame, text="много").pack(side="left")
        self.sens_label = ctk.CTkLabel(sens_frame, text=f"порог: {_sensitivity_to_threshold(5)}", width=90,
                                       text_color="gray")
        self.sens_label.pack(side="left", padx=8)

        self.process_btn = ctk.CTkButton(self, text="Обработать", height=40,
                                         font=ctk.CTkFont(size=15, weight="bold"),
                                         command=self.start_processing, state="disabled")
        self.process_btn.pack(padx=20, pady=10)

        self.progress = ctk.CTkProgressBar(self, width=660)
        self.progress.pack(padx=20, pady=(0, 4))
        self.progress.set(0)

        self.status_label = ctk.CTkLabel(self, text="", text_color="gray")
        self.status_label.pack(padx=20, pady=2)

        self.log_box = ctk.CTkTextbox(self, height=220, font=ctk.CTkFont(family="Courier New", size=12))
        self.log_box.pack(fill="both", expand=True, padx=20, pady=(4, 20))

    def _on_slider(self, value):
        self.sens_label.configure(text=f"порог: {_sensitivity_to_threshold(float(value))}")

    # ── Действия ────────────────────────────────────────────────────────────

    def select_video(self):
        path = filedialog.askopenfilename(
            title="Выберите видео",
            filetypes=[("Видео файлы", "*.mp4 *.mkv *.avi *.mov *.webm"), ("Все файлы", "*.*")]
        )
        if path:
            self.video_path = path
            self.path_label.configure(text=os.path.basename(path))
            self.process_btn.configure(state="normal")

    def start_processing(self):
        self.process_btn.configure(state="disabled")
        self.log_box.delete("1.0", "end")
        self.progress.set(0)
        self.status_label.configure(text="")
        threading.Thread(target=self._worker, daemon=True).start()

    # ── Воркер ──────────────────────────────────────────────────────────────

    def _worker(self):
        try:
            video_path = self.video_path
            video_name = os.path.splitext(os.path.basename(video_path))[0].strip()
            out_dir = os.path.join(os.path.dirname(video_path), video_name)
            os.makedirs(out_dir, exist_ok=True)

            # ── Шаг 1: транскрибация ─────────────────────────────────────
            self._ui(lambda: self.status_label.configure(text="Шаг 1/3 — Транскрибация..."))
            model_name = WHISPER_MODELS[self.model_var.get()]
            self._log(f"Загрузка модели Whisper «{model_name}»...")
            model = whisper.load_model(model_name)
            self._log("Транскрибация видео (может занять несколько минут)...")
            result = model.transcribe(video_path, verbose=None)

            # Сохраняем полный текст
            txt_path = os.path.join(out_dir, f"{video_name}.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(result["text"].strip())
            self._log(f"✓ Транскрипция сохранена: {os.path.basename(txt_path)}")

            # Извлекаем сегменты с таймингами
            segments = [
                {"start": round(s["start"], 2), "end": round(s["end"], 2), "text": s["text"]}
                for s in result.get("segments", [])
            ]
            self._log(f"  Сегментов с таймингами: {len(segments)}")
            self._ui(lambda: self.progress.set(0.4))

            # ── Шаг 2: скриншоты ─────────────────────────────────────────
            self._ui(lambda: self.status_label.configure(text="Шаг 2/3 — Скриншоты..."))
            threshold = _sensitivity_to_threshold(self.sens_slider.get())
            self._log(f"Анализ смены кадров (порог {threshold})...")

            cap = cv2.VideoCapture(video_path)
            fps = cap.get(cv2.CAP_PROP_FPS) or 25
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

            scene_frames = self._detect_scenes(video_path, threshold)
            if not scene_frames:
                self._log(f"Сцены не найдены — скриншот каждые {FALLBACK_INTERVAL_SEC} с.")
                interval = int(fps * FALLBACK_INTERVAL_SEC)
                scene_frames = list(range(0, total_frames, interval))

            self._log(f"Найдено сцен: {len(scene_frames)}")

            screenshots = []
            for i, frame_num in enumerate(scene_frames, 1):
                time_sec = round(frame_num / fps, 2)
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_num)
                ret, frame = cap.read()
                if ret:
                    filename = f"{i}.jpg"
                    cv2.imwrite(os.path.join(out_dir, filename), frame)
                    screenshots.append({"index": i, "time_sec": time_sec, "filename": filename})
                self._ui(lambda p=0.4 + 0.3 * (i / len(scene_frames)): self.progress.set(p))

            cap.release()
            self._log(f"✓ Скриншотов сохранено: {len(screenshots)}")

            # Сохраняем манифест
            manifest = {"video": os.path.basename(video_path), "segments": segments, "screenshots": screenshots}
            with open(os.path.join(out_dir, "manifest.json"), "w", encoding="utf-8") as f:
                json.dump(manifest, f, ensure_ascii=False, indent=2)
            self._log("✓ Манифест сохранён: manifest.json")
            self._ui(lambda: self.progress.set(0.7))

            # ── Шаг 3: Word-документ ─────────────────────────────────────
            self._ui(lambda: self.status_label.configure(text="Шаг 3/3 — Создание Word..."))
            self._log("Генерация Word-документа...")
            docx_path = _generate_word(out_dir, video_name, segments, screenshots)
            self._log(f"✓ Word сохранён: {os.path.basename(docx_path)}")
            self._ui(lambda: self.progress.set(1.0))

            # ── Готово ───────────────────────────────────────────────────
            self._ui(lambda: self.status_label.configure(text="Готово!", text_color="green"))
            self._ui(lambda: messagebox.showinfo(
                "Готово",
                f"Обработка завершена!\n\n"
                f"Транскрипция: {os.path.basename(txt_path)}\n"
                f"Скриншотов: {len(screenshots)}\n"
                f"Word-документ: {os.path.basename(docx_path)}\n\n"
                f"Папка с результатами:\n{out_dir}"
            ))

        except Exception as e:
            self._log(f"ОШИБКА: {e}")
            self._ui(lambda: self.status_label.configure(text="Ошибка", text_color="red"))
        finally:
            self._ui(lambda: self.process_btn.configure(state="normal"))

    def _detect_scenes(self, video_path: str, threshold: float) -> list[int]:
        video = open_video(video_path)
        scene_manager = SceneManager()
        scene_manager.add_detector(ContentDetector(threshold=threshold))
        scene_manager.detect_scenes(video, show_progress=False)
        return [s[0].get_frames() for s in scene_manager.get_scene_list()]

    def _log(self, msg: str):
        self.after(0, lambda m=msg: (self.log_box.insert("end", m + "\n"), self.log_box.see("end")))

    def _ui(self, fn):
        self.after(0, fn)


if __name__ == "__main__":
    app = App()
    app.mainloop()
